from __future__ import annotations

import csv
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSET_DIR = DOCS_DIR / "assets" / "tutorial"
CSV_PATH = ROOT / "data" / "demo_data_analysis_tool.csv"
OUTPUT_PATH = DOCS_DIR / "data-analysis-tool-user-tutorial.docx"

FONT_CN = "Microsoft YaHei"
FONT_MONO = "Consolas"
BLUE = "1E4E8C"
TEAL = "168197"
LIGHT_BLUE = "EAF3F8"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "6B7280"
DARK = "111827"


def set_run_font(run, size=None, bold=None, color=None, font=FONT_CN):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.08):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=DARK, size=9.2, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    set_paragraph_spacing(p, after=0, line=1.05)
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_padding(cell, top=90, left=120, bottom=90, right=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for edge, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2EC", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, width_pct=100):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), str(width_pct * 50))


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    if level == 1:
        paragraph.style = "Heading 1"
        set_paragraph_spacing(paragraph, before=8, after=8, line=1.05)
        size, color = 18, BLUE
    elif level == 2:
        paragraph.style = "Heading 2"
        set_paragraph_spacing(paragraph, before=7, after=5, line=1.05)
        size, color = 13, TEAL
    else:
        paragraph.style = "Heading 3"
        set_paragraph_spacing(paragraph, before=5, after=4, line=1.05)
        size, color = 11.5, DARK
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)
    return paragraph


def add_body(doc, text, after=5):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=after, line=1.12)
    run = paragraph.add_run(text)
    set_run_font(run, size=10, color=DARK)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(paragraph, after=3, line=1.08)
        run = paragraph.add_run(item)
        set_run_font(run, size=9.8, color=DARK)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        set_paragraph_spacing(paragraph, after=3, line=1.08)
        run = paragraph.add_run(item)
        set_run_font(run, size=9.8, color=DARK)


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 100)
    set_table_borders(table, color="D1D5DB", size="4")
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_GRAY)
    set_cell_padding(cell, top=110, left=150, bottom=110, right=150)
    cell.text = ""
    for line_no, line in enumerate(code.strip("\n").splitlines()):
        p = cell.paragraphs[0] if line_no == 0 else cell.add_paragraph()
        set_paragraph_spacing(p, after=0, line=1.0)
        run = p.add_run(line)
        set_run_font(run, size=8.5, color="111827", font=FONT_MONO)
    doc.add_paragraph()


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=2, after=8, line=1.0)
    run = paragraph.add_run(text)
    set_run_font(run, size=8.5, color=MID_GRAY)


def add_image(doc, filename, caption):
    path = ASSET_DIR / filename
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=4, after=2, line=1.0)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.45))
    add_caption(doc, caption)


def add_info_box(doc, title, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 100)
    set_table_borders(table, color="C7D7E5", size="6")
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_BLUE)
    set_cell_padding(cell, top=130, left=160, bottom=130, right=160)
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=3, line=1.05)
    run = p.add_run(title)
    set_run_font(run, size=10.5, bold=True, color=BLUE)
    for line in lines:
        p = cell.add_paragraph()
        set_paragraph_spacing(p, after=2, line=1.05)
        run = p.add_run(line)
        set_run_font(run, size=9.2, color=DARK)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table, 100)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        shade_cell(cell, BLUE)
        set_cell_padding(cell)
        set_cell_text(cell, text, bold=True, color="FFFFFF", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        if widths:
            cell.width = Inches(widths[i])

    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            set_cell_padding(cell)
            set_cell_text(cell, text, size=8.8)
            if i == 0:
                shade_cell(cell, "F8FAFC")
            if widths:
                cell.width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_cover(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(17)
    section.right_margin = Mm(17)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, before=35, after=10, line=1.0)
    run = paragraph.add_run("数据分析工具")
    set_run_font(run, size=30, bold=True, color=BLUE)

    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=22, line=1.0)
    run = paragraph.add_run("使用教程 Word 版")
    set_run_font(run, size=22, bold=True, color=TEAL)

    add_info_box(
        doc,
        "文档用途",
        [
            "用于产品演示、客户培训和内部交付说明。",
            "覆盖上传、表格、筛选、图表、统计、相关性、分组、SQL、JOIN、关系图、时间对比和导出。",
            f"生成日期：{date.today().isoformat()}",
        ],
    )

    add_table(
        doc,
        ["项目", "文件"],
        [
            ["演示数据", "data/demo_data_analysis_tool.csv"],
            ["能力清单", "docs/DATA_ANALYSIS_TOOLS_AND_CAPABILITIES.md"],
            ["截图目录", "docs/assets/tutorial/"],
        ],
        widths=[1.6, 4.8],
    )
    doc.add_page_break()


def add_toc(doc):
    add_heading(doc, "目录", 1)
    rows = [
        ["1", "准备工作", "启动服务，确认演示数据和截图资源。"],
        ["2", "上传演示数据", "上传 CSV，确认 64 行 33 列。"],
        ["3", "主工作区", "认识文件信息、数据源、筛选器和标签页。"],
        ["4", "表格与筛选", "按列查看、排序、隐藏字段和组合筛选。"],
        ["5", "日期和时间分布", "使用日期范围筛选，观察月度趋势。"],
        ["6", "图表、统计、相关性", "展示业务指标、统计摘要和相关矩阵。"],
        ["7", "分组、SQL、JOIN", "完成聚合查询、SQL 查询和结果合并。"],
        ["8", "时间对比和导出", "检查日期顺序异常，导出 CSV。"],
        ["9", "演示路线和字段说明", "按固定路线完成培训讲解。"],
    ]
    add_table(doc, ["序号", "章节", "说明"], rows, widths=[0.55, 1.8, 4.2])


def add_content(doc):
    add_heading(doc, "1. 准备工作", 1)
    add_body(doc, "在项目根目录启动本地服务。")
    add_code_block(doc, "npm run dev")
    add_body(doc, "打开终端输出的本地地址。常见地址如下。")
    add_code_block(doc, "http://127.0.0.1:5002/")
    add_info_box(
        doc,
        "检查项",
        [
            "确认演示数据存在：data/demo_data_analysis_tool.csv。",
            "确认截图目录存在：docs/assets/tutorial/。",
            "确认浏览器能打开本地页面。",
        ],
    )

    add_heading(doc, "2. 上传演示数据", 1)
    add_body(doc, "进入首页后，点击“选择文件”，选择演示 CSV。")
    add_image(doc, "01-upload-empty.png", "图 1：上传入口")
    add_body(doc, "上传成功后，页面会显示“已加载 64 行和 33 列”。")
    add_table(
        doc,
        ["动作", "结果"],
        [
            ["选择 CSV 文件", "系统读取字段和数据行。"],
            ["加载完成", "出现数据预览、筛选器和分析标签页。"],
            ["开始演示", "先看表格，再逐步进入图表、统计和 SQL。"],
        ],
        widths=[2.0, 4.4],
    )

    add_heading(doc, "3. 认识主工作区", 1)
    add_image(doc, "10-overview-after-upload.png", "图 2：上传后的主工作区")
    add_table(
        doc,
        ["区域", "用途"],
        [
            ["当前文件", "查看文件名、行数和列数。"],
            ["数据源", "切换主数据、SQL 结果、分组结果和 JOIN 结果。"],
            ["筛选器", "按文本、数字和日期缩小数据范围。"],
            ["日期范围筛选", "用日期列限制分析窗口。"],
            ["时间分布", "查看记录按时间分布和累计趋势。"],
            ["标签页", "切换表格、图表、统计、相关性、分组、时间和 SQL。"],
        ],
        widths=[1.5, 4.9],
    )

    add_heading(doc, "4. 表格与筛选", 1)
    add_body(doc, "表格页适合先核对数据质量。先检查日期、收入、利润、满意度和退货字段。")
    add_table(
        doc,
        ["演示目标", "字段", "操作符", "值"],
        [
            ["查看东部区域", "region", "等于", "East"],
            ["查看高收入订单", "revenue", "大于", "5000"],
            ["查看退货订单", "returned", "等于", "Yes"],
        ],
        widths=[1.8, 1.6, 1.2, 1.8],
    )
    add_body(doc, "多个筛选条件会按 AND 生效。")
    add_body(doc, "这不是为了增加配置复杂度，而是为了让业务人员逐步缩小问题范围。")

    add_heading(doc, "5. 日期范围和时间分布", 1)
    add_body(doc, "在日期范围筛选区选择 order_date，然后拖动开始和结束日期。")
    add_table(
        doc,
        ["演示目标", "操作"],
        [
            ["只看 2025 年订单", "把开始日期调整到 2025-01-01 附近。"],
            ["看活动周期", "结合 campaign 字段和日期范围筛选。"],
            ["看业务节奏", "应用筛选后观察时间分布图变化。"],
        ],
        widths=[2.0, 4.4],
    )
    add_body(doc, "演示数据跨越 2024-01 到 2025-12。默认时间粒度通常是 Month。")

    add_heading(doc, "6. 图表页", 1)
    add_image(doc, "11-charts-tab.png", "图 3：图表标签页")
    add_table(
        doc,
        ["图表", "适合场景"],
        [
            ["柱状图", "比较不同记录的数值大小。"],
            ["折线图", "查看指标变化趋势。"],
            ["饼图", "查看前 20 行的占比结构。"],
        ],
        widths=[1.5, 4.9],
    )
    add_body(doc, "推荐演示 revenue、profit、units、marketing_spend 和 satisfaction_score。")

    add_heading(doc, "7. 统计页", 1)
    add_image(doc, "12-statistics-tab.png", "图 4：统计标签页")
    add_table(
        doc,
        ["指标", "含义"],
        [
            ["数量", "有效数字值数量。"],
            ["平均值", "数值平均水平。"],
            ["中位数", "中间位置的数值。"],
            ["总和", "全部数值之和。"],
            ["最小值 / 最大值", "识别极端值和范围。"],
        ],
        widths=[1.7, 4.7],
    )

    add_heading(doc, "8. 相关性和散点回归", 1)
    add_image(doc, "13-correlation-tab.png", "图 5：相关性标签页")
    add_table(
        doc,
        ["字段组合", "预期现象"],
        [
            ["satisfaction_score 与 renewal_likelihood", "强正相关。"],
            ["revenue 与 marketing_spend", "强正相关。"],
            ["revenue 与 profit", "强正相关。"],
            ["support_tickets 与 satisfaction_score", "负相关。"],
        ],
        widths=[3.1, 3.3],
    )
    add_body(doc, "继续向下可以查看散点图和回归线。推荐 X 轴选择 revenue，Y 轴选择 profit。")

    add_heading(doc, "9. 分组与聚合", 1)
    add_image(doc, "14-groupby-tab.png", "图 6：分组标签页")
    add_table(
        doc,
        ["演示", "分组列", "聚合函数", "聚合列", "输出列名"],
        [
            ["按区域收入", "region", "sum", "revenue", "total_revenue"],
            ["按品类利润", "category", "avg", "profit", "avg_profit"],
            ["客户数量", "region", "countDistinct", "customer_id", "unique_customers"],
        ],
        widths=[1.35, 1.25, 1.25, 1.25, 1.3],
    )
    add_body(doc, "执行后，结果会进入查询结果列表。分组结果也可以继续进入图表和统计页。")

    add_heading(doc, "10. SQL 查询", 1)
    add_image(doc, "16-sql-tab.png", "图 7：SQL 标签页")
    add_body(doc, "SQL 里当前数据集用 ? 表示。先跑汇总查询。")
    add_code_block(
        doc,
        """
SELECT
  region,
  category,
  SUM(revenue) AS total_revenue,
  AVG(profit) AS avg_profit,
  COUNT(*) AS orders
FROM ?
GROUP BY region, category
ORDER BY total_revenue DESC
""",
    )
    add_body(doc, "再跑异常订单查询。")
    add_code_block(
        doc,
        """
SELECT
  order_id,
  customer_name,
  order_date,
  ship_date,
  returned,
  support_tickets,
  satisfaction_score
FROM ?
WHERE returned = 'Yes'
ORDER BY support_tickets DESC
""",
    )

    add_heading(doc, "11. JOIN 和关系图", 1)
    add_body(doc, "JOIN 需要至少两个已保存查询结果。先生成客户维度表。")
    add_code_block(
        doc,
        """
SELECT DISTINCT
  customer_id,
  customer_name,
  segment,
  region
FROM ?
""",
    )
    add_body(doc, "再生成订单事实表。")
    add_code_block(
        doc,
        """
SELECT
  order_id,
  customer_id,
  order_date,
  product_name,
  revenue,
  profit
FROM ?
""",
    )
    add_table(
        doc,
        ["设置项", "值"],
        [
            ["左表", "客户维度表"],
            ["右表", "订单事实表"],
            ["左关联列", "customer_id"],
            ["右关联列", "customer_id"],
            ["JOIN 类型", "INNER JOIN"],
        ],
        widths=[2.0, 4.4],
    )
    add_body(doc, "执行 JOIN 后，SQL 页底部会显示关系图。白色节点代表原始查询结果，蓝色节点代表 JOIN 结果。")

    add_heading(doc, "12. 时间对比", 1)
    add_image(doc, "15-time-compare-tab.png", "图 8：时间对比标签页")
    add_table(
        doc,
        ["设置项", "值"],
        [
            ["列 A（预期更早）", "order_date"],
            ["列 B（预期更晚）", "ship_date"],
        ],
        widths=[2.3, 4.1],
    )
    add_table(
        doc,
        ["记录", "异常"],
        [
            ["ORD-1008", "ship_date 早于 order_date。"],
            ["ORD-1052", "ship_date 早于 order_date。"],
        ],
        widths=[1.8, 4.6],
    )
    add_body(doc, "时间对比会把结果分成 A -> B、B -> A、A = B 和缺失四类。")

    doc.add_page_break()
    add_heading(doc, "13. 导出结果", 1)
    add_table(
        doc,
        ["导出位置", "导出内容"],
        [
            ["顶部“导出当前数据”", "当前数据源。"],
            ["查询结果卡片导出按钮", "SQL、分组、JOIN 结果。"],
            ["时间对比页导出按钮", "当前时间对比筛选结果。"],
        ],
        widths=[2.4, 4.0],
    )
    add_body(doc, "导出格式是 CSV。导出后可以发给同事，也可以继续用 Excel 处理。")

    add_heading(doc, "14. 推荐演示流程", 1)
    add_numbered(
        doc,
        [
            "上传 data/demo_data_analysis_tool.csv。",
            "介绍主工作区和数据源。",
            "进入表格页，说明字段类型和列选择。",
            "添加筛选：region = East。",
            "添加筛选：revenue > 5000。",
            "选择 order_date 做日期范围筛选。",
            "看时间分布图。",
            "切到图表页，选择 revenue。",
            "切到统计页，解释收入、利润、满意度。",
            "切到相关性页，解释强相关关系。",
            "切到分组页，按 region 汇总 revenue。",
            "切到 SQL 页，运行分组 SQL。",
            "生成两个 SQL 结果，演示 JOIN。",
            "查看关系图。",
            "切到时间页，演示 order_date 和 ship_date 顺序异常。",
            "导出当前结果。",
        ],
    )

    add_heading(doc, "15. 常见问题", 1)
    add_table(
        doc,
        ["问题", "处理方式"],
        [
            ["上传后没有统计", "检查数据里是否有数值列。"],
            ["图表只显示部分数据", "基础图表默认展示前 20 行。"],
            ["分组图没有出现", "先执行一次分组聚合。"],
            ["JOIN 面板不可用", "先保存至少两个 SQL 或分组结果。"],
            ["时间页不可用", "数据集至少需要两个日期列。"],
            ["导出的 CSV 打开乱码", "用支持 UTF-8 的表格软件打开。"],
        ],
        widths=[2.2, 4.2],
    )

    add_heading(doc, "16. 演示数据字段说明", 1)
    with CSV_PATH.open(newline="", encoding="utf-8") as fh:
        reader = csv.reader(fh)
        headers = next(reader)
    field_rows = []
    descriptions = {
        "order_id": "订单编号",
        "customer_id": "客户编号，支持 JOIN",
        "customer_name": "客户名称",
        "segment": "客户分层",
        "region": "区域分组",
        "product_name": "产品名称",
        "category": "产品品类",
        "order_date": "下单日期",
        "ship_date": "发货日期",
        "delivery_date": "送达日期",
        "invoice_date": "开票日期",
        "payment_date": "付款日期",
        "units": "数量",
        "unit_price": "单价",
        "discount_pct": "折扣比例",
        "revenue": "收入",
        "cost": "成本",
        "profit": "利润",
        "marketing_spend": "市场费用",
        "satisfaction_score": "满意度",
        "support_tickets": "支持工单数",
        "renewal_likelihood": "续约可能性",
        "returned": "是否退货",
        "warehouse_id": "仓库编号",
        "supplier_id": "供应商编号",
        "campaign": "营销活动",
    }
    for field in headers:
        field_rows.append([field, descriptions.get(field, "演示分析字段")])
    add_table(doc, ["字段", "用途"], field_rows, widths=[2.35, 4.05])


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10)

    for style_name in ("Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, after=0, line=1.0)
        run = paragraph.add_run("数据分析工具使用教程")
        set_run_font(run, size=8, color=MID_GRAY)


def build():
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_toc(doc)
    add_content(doc)
    add_footer(doc)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build()
